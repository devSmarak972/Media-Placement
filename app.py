import os
import io
import requests
import zipfile
import shutil
from datetime import datetime
from flask import Flask, render_template, redirect, url_for, flash, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_wtf.csrf import CSRFProtect
import logging
import pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image
from bs4 import BeautifulSoup
import tempfile

from config import Config
from models import db, MediaPlacement, GoogleCredential
from forms import AddPlacementForm, GoogleCredentialForm
from utils import setup_logging, take_screenshot
from google_integration import google_bp, docket_bp, get_google_docs_content, get_google_sheets_content
from parsers import extract_links, parse_media_links

# Initialize Flask app
app = Flask(__name__)
app.config.from_object(Config)

# Setup logging
setup_logging(app)

# Initialize extensions
db.init_app(app)
migrate = Migrate(app, db)
csrf = CSRFProtect(app)

# Add utility functions to template context
@app.context_processor
def utility_processor():
    return {
        'now': datetime.utcnow
    }

# Register blueprints
app.register_blueprint(google_bp)
app.register_blueprint(docket_bp)

@app.route('/')
def index():
    return redirect(url_for('dashboard'))

@app.route('/dashboard')
def dashboard():
    placements = MediaPlacement.query.order_by(MediaPlacement.created_at.desc()).all()
    return render_template('dashboard.html', placements=placements)

@app.route('/add_placement', methods=['GET', 'POST'])
def add_placement():
    form = AddPlacementForm()
    if form.validate_on_submit():
        # Handle direct text input
        if form.input_type.data == 'direct':
            text_content = form.text_input.data
            links = extract_links(text_content)
            if not links:
                flash('No valid media links found in the provided text.', 'warning')
                return render_template('add_placement.html', form=form)
                
            # Parse and save links
            added_count = 0
            for link in links:
                placement_data = parse_media_links(link)
                if placement_data:
                    placement = MediaPlacement(
                        url=link,
                        title=placement_data.get('title', ''),
                        source=placement_data.get('source', ''),
                        publication_date=placement_data.get('date'),
                        media_type=placement_data.get('type', 'article')
                    )
                    db.session.add(placement)
                    added_count += 1
            
            if added_count > 0:
                try:
                    db.session.commit()
                    flash(f'Successfully added {added_count} media placements!', 'success')
                    return redirect(url_for('dashboard'))
                except Exception as e:
                    db.session.rollback()
                    app.logger.error(f"Database error when adding placements: {str(e)}")
                    flash(f'Error saving to database: {str(e)}', 'danger')
                    return render_template('add_placement.html', form=form)
            else:
                flash('Could not extract media information from the provided links.', 'warning')
        
        # Handle Google Docs
        elif form.input_type.data == 'gdoc':
            doc_id = form.google_doc_id.data
            print("Google Doc ID:", doc_id)
            try:
                # Get the first Google credential (since we no longer have user-specific credentials)
                google_cred = GoogleCredential.query.first()
                print(google_cred)
                if not google_cred:
                    flash('Google API credentials are not set up. Please add them in Settings.', 'warning')
                    return redirect(url_for('settings'))
                    
                content = get_google_docs_content(doc_id)
                links = extract_links(content)
                print("LInks from Google Doc:", links)
                if not links:
                    flash('No valid media links found in the Google Doc.', 'warning')
                    return render_template('add_placement.html', form=form)
                
                # Parse and save links
                added_count = 0
                for link in links:
                    placement_data = parse_media_links(link)
                    if placement_data:
                        placement = MediaPlacement(
                            url=link,
                            title=placement_data.get('title', ''),
                            source=placement_data.get('source', ''),
                            publication_date=placement_data.get('date'),
                            media_type=placement_data.get('type', 'article')
                        )
                        db.session.add(placement)
                        added_count += 1
                
                if added_count > 0:
                    try:
                        db.session.commit()
                        flash(f'Successfully added {added_count} media placements from Google Doc!', 'success')
                        return redirect(url_for('dashboard'))
                    except Exception as e:
                        db.session.rollback()
                        app.logger.error(f"Database error when adding placements from Google Doc: {str(e)}")
                        flash(f'Error saving to database: {str(e)}', 'danger')
                        return render_template('add_placement.html', form=form)
                else:
                    flash('Could not extract media information from the links in the Google Doc.', 'warning')
            except Exception as e:
                app.logger.error(f"Error accessing Google Doc: {e}")
                flash(f'Error accessing Google Doc: {str(e)}', 'danger')
        
        # Handle Google Sheets
        elif form.input_type.data == 'gsheet':
            sheet_id = form.google_sheet_id.data
            try:
                # Get the first Google credential (since we no longer have user-specific credentials)
                google_cred = GoogleCredential.query.first()
                if not google_cred:
                    flash('Google API credentials are not set up. Please add them in Settings.', 'warning')
                    return redirect(url_for('settings'))
                    
                content = get_google_sheets_content(sheet_id)
                links = extract_links(content)
                
                if not links:
                    flash('No valid media links found in the Google Sheet.', 'warning')
                    return render_template('add_placement.html', form=form)
                
                # Parse and save links
                added_count = 0
                for link in links:
                    placement_data = parse_media_links(link)
                    if placement_data:
                        placement = MediaPlacement(
                            url=link,
                            title=placement_data.get('title', ''),
                            source=placement_data.get('source', ''),
                            publication_date=placement_data.get('date'),
                            media_type=placement_data.get('type', 'article')
                        )
                        db.session.add(placement)
                        added_count += 1
                
                if added_count > 0:
                    try:
                        db.session.commit()
                        flash(f'Successfully added {added_count} media placements from Google Sheet!', 'success')
                        return redirect(url_for('dashboard'))
                    except Exception as e:
                        db.session.rollback()
                        app.logger.error(f"Database error when adding placements from Google Sheet: {str(e)}")
                        flash(f'Error saving to database: {str(e)}', 'danger')
                        return render_template('add_placement.html', form=form)
                else:
                    flash('Could not extract media information from the links in the Google Sheet.', 'warning')
            except Exception as e:
                app.logger.error(f"Error accessing Google Sheet: {e}")
                flash(f'Error accessing Google Sheet: {str(e)}', 'danger')
        
    return render_template('add_placement.html', form=form)

@app.route('/placement/<int:placement_id>')
def view_placement(placement_id):
    placement = MediaPlacement.query.filter_by(id=placement_id).first_or_404()
    return render_template('view_placement.html', placement=placement)

@app.route('/placement/<int:placement_id>/delete', methods=['POST'])
def delete_placement(placement_id):
    placement = MediaPlacement.query.filter_by(id=placement_id).first_or_404()
    try:
        db.session.delete(placement)
        db.session.commit()
        flash('Media placement deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Database error when deleting placement: {str(e)}")
        flash(f'Error deleting placement: {str(e)}', 'danger')
    return redirect(url_for('dashboard'))

@app.route('/settings', methods=['GET', 'POST'])
def settings():
    google_form = GoogleCredentialForm()
    
    # Check if we already have Google credentials
    google_cred = GoogleCredential.query.first()
    
    if google_form.validate_on_submit():
        try:
            if google_cred:
                google_cred.api_key = google_form.api_key.data
                flash('Google API Key updated successfully!', 'success')
            else:
                google_cred = GoogleCredential(
                    api_key=google_form.api_key.data
                )
                db.session.add(google_cred)
                flash('Google API Key added successfully!', 'success')
            
            db.session.commit()
            return redirect(url_for('settings'))
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Database error when saving Google credentials: {str(e)}")
            flash(f'Error saving credentials: {str(e)}', 'danger')
    
    # Pre-populate form if credentials exist
    if google_cred and request.method == 'GET':
        google_form.api_key.data = google_cred.api_key
        
    return render_template('settings.html', google_form=google_form)

@app.errorhandler(404)
def not_found_error(error):
    return render_template('error.html', error='Page not found'), 404

@app.route('/export/excel')
def export_excel():
    """Export all media placements to an Excel file directly."""
    try:
        # Get all placements
        placements = MediaPlacement.query.all()
        
        if not placements:
            flash('No media placements found to export.', 'info')
            return redirect(url_for('dashboard'))
        
        # Prepare data for Excel
        data = []
        for placement in placements:
            data.append({
                'Title': placement.title or "Untitled",
                'URL': placement.url,
                'Source': placement.source or "Unknown",
                'Publication Date': str(placement.publication_date) if placement.publication_date else "Unknown",
                'Media Type': placement.media_type,
                'Docket Link': placement.docket_url or "No docket",
                'Created': placement.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'Updated': placement.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                'Notes': placement.notes or ""
            })
        
        # Create a Pandas DataFrame
        df = pd.DataFrame(data)
        
        # Create an Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Media Placements', index=False)
            
            # Auto-adjust columns' width
            worksheet = writer.sheets['Media Placements']
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = max_length
        
        output.seek(0)
        
        # Create timestamp for filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Return the Excel file as a download
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'media_placements_{timestamp}.xlsx'
        )
        
    except Exception as e:
        app.logger.error(f"Error exporting to Excel: {str(e)}")
        flash(f'Error exporting to Excel: {str(e)}', 'danger')
        return redirect(url_for('dashboard'))

@app.route('/create-docx-docket/<int:placement_id>')
def create_docx_docket(placement_id):
    """Create a Word document docket for a specific media placement."""
    # Get the placement
    placement = MediaPlacement.query.filter_by(id=placement_id).first_or_404()
    
    # # Show loading screen first
    # if request.args.get('start') != 'true':
    #     return render_template(
    #         'loading.html',
    #         action_title='Creating Docket',
    #         message=f'Taking a screenshot and preparing docket for "{placement.title or "Untitled"}"'
    #     )
    
    try:
        
        # Create a new Word document
        doc = Document()
        
        # Add title with formatting
        doc.add_heading(placement.title or "Untitled Article", level=1)
        
        # Add basic info section
        doc.add_heading("Basic Information", level=2)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Add basic info rows
        cells = table.rows[0].cells
        cells[0].text = "URL"
        cells[1].text = placement.url
        
        cells = table.rows[1].cells
        cells[0].text = "Source"
        cells[1].text = placement.source or "Unknown"
        
        cells = table.rows[2].cells
        cells[0].text = "Publication Date"
        cells[1].text = str(placement.publication_date) if placement.publication_date else "Unknown"
        
        cells = table.rows[3].cells
        cells[0].text = "Media Type"
        cells[1].text = placement.media_type.title()
        
        cells = table.rows[4].cells
        cells[0].text = "Created"
        cells[1].text = placement.created_at.strftime('%Y-%m-%d %H:%M:%S')
        
        # Add screenshot section
        doc.add_heading("Screenshot", level=2)
        
        # Create a loading paragraph
        print("Now taking screenshot...")
        # Try to take a screenshot of the URL using Selenium
        try:
            # Create a temporary file for the screenshot
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_screenshot_path = temp_file.name
            
            # Show status in logs
            app.logger.info(f"Taking screenshot of {placement.url}")
            
            # Take the screenshot using our utility function
            screenshot_path = take_screenshot(placement.url, temp_screenshot_path)
            
            if screenshot_path:
                # Add the screenshot to the document
                doc.add_picture(screenshot_path, width=Inches(6.0))
                doc.add_paragraph(f"Screenshot taken on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                # Clean up the temporary file
                try:
                    os.unlink(temp_screenshot_path)
                except:
                    pass
            else:
                doc.add_paragraph("Screenshot could not be captured. The website may be protected or requires authentication.")
        except Exception as e:
            app.logger.error(f"Error taking screenshot for {placement.url}: {str(e)}")
            doc.add_paragraph(f"Error capturing screenshot: {str(e)}")
        
        # Add summary section
        doc.add_heading("Summary", level=2)
        print("Now extracting summary...")
        # Try to extract a summary
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(placement.url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Try to get the main article content
            article = soup.find('article') or soup.find(class_=['article', 'post', 'content', 'main-content'])
            
            if article:
                paragraphs = article.find_all('p')
                content = ' '.join([p.get_text().strip() for p in paragraphs[:5]])  # Get first 5 paragraphs
            else:
                # Fallback to all paragraphs
                paragraphs = soup.find_all('p')
                content = ' '.join([p.get_text().strip() for p in paragraphs[:5]])
            
            # Limit summary length
            if content:
                content = content[:1000] + '...' if len(content) > 1000 else content
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("No text content could be extracted from this page.")
                
        except Exception as e:
            app.logger.error(f"Error extracting summary for {placement.url}: {str(e)}")
            doc.add_paragraph(f"Error extracting summary: {str(e)}")
        
        # Add notes section
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(placement.notes or "No notes available")
        
        # Save the document to a BytesIO object
        docx_file = io.BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        # Create a sanitized title for filename
        safe_title = ''.join(c for c in (placement.title or "untitled") if c.isalnum() or c in ' -_')[:30]
        safe_title = safe_title.replace(' ', '_')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Return the document as a download
        return send_file(
            docx_file,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'docket_{safe_title}_{timestamp}.docx'
        )
        
    except Exception as e:
        app.logger.error(f"Error creating DOCX docket: {str(e)}")
        flash(f'Error creating DOCX docket: {str(e)}', 'danger')
        return redirect(url_for('view_placement', placement_id=placement_id))

@app.route('/export/excel/<int:placement_id>')
def export_single_excel(placement_id):
    """Export a single media placement to an Excel file directly."""
    try:
        # Get the specific placement
        placement = MediaPlacement.query.filter_by(id=placement_id).first_or_404()
        
        # Prepare data for Excel
        data = [{
            'Title': placement.title or "Untitled",
            'URL': placement.url,
            'Source': placement.source or "Unknown",
            'Publication Date': str(placement.publication_date) if placement.publication_date else "Unknown",
            'Media Type': placement.media_type,
            'Docket Link': placement.docket_url or "No docket",
            'Created': placement.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'Updated': placement.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            'Notes': placement.notes or ""
        }]
        
        # Create a Pandas DataFrame
        df = pd.DataFrame(data)
        
        # Create an Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Media Placement', index=False)
            
            # Auto-adjust columns' width
            worksheet = writer.sheets['Media Placement']
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = max_length
        
        output.seek(0)
        
        # Create a sanitized title for filename
        safe_title = ''.join(c for c in (placement.title or "untitled") if c.isalnum() or c in ' -_')[:30]
        safe_title = safe_title.replace(' ', '_')
        
        # Return the Excel file as a download
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'media_placement_{safe_title}_{placement.id}.xlsx'
        )
        
    except Exception as e:
        app.logger.error(f"Error exporting to Excel: {str(e)}")
        flash(f'Error exporting to Excel: {str(e)}', 'danger')
        return redirect(url_for('view_placement', placement_id=placement_id))

@app.route('/export/complete')
def export_complete_package():
    """Export all media placements with their dockets as a complete ZIP package."""
    try:
        # Create a loading page first
        # if request.args.get('start') != 'true':
        #     return render_template(
        #         'loading.html',
        #         action_title='Creating Complete Export Package',
        #         message='Generating dockets and preparing ZIP file with all data...'
        #     )
            
        # Get all placements
        placements = MediaPlacement.query.all()
        
        if not placements:
            flash('No media placements found to export.', 'info')
            return redirect(url_for('dashboard'))
        
        # Create temporary directory to store files
        temp_dir = tempfile.mkdtemp()
        dockets_dir = os.path.join(temp_dir, 'dockets')
        os.makedirs(dockets_dir, exist_ok=True)
        
        # Prepare data for Excel with local hyperlinks to dockets
        data = []
        
        # Process each placement and create docket if needed
        for placement in placements:
            # Create a unique filename for the docket
            safe_title = ''.join(c for c in (placement.title or "untitled") if c.isalnum() or c in ' -_')[:30]
            safe_title = safe_title.replace(' ', '_')
            docket_filename = f'docket_{placement.id}_{safe_title}.docx'
            docket_path = os.path.join(dockets_dir, docket_filename)
            
            # Create the docket file
            create_docket_for_export(placement, docket_path)
            
            # Add data with hyperlink to local docket file
            docket_rel_path = f'dockets/{docket_filename}'
            data.append({
                'ID': placement.id,
                'Title': placement.title or "Untitled",
                'URL': placement.url,
                'Source': placement.source or "Unknown",
                'Publication Date': str(placement.publication_date) if placement.publication_date else "Unknown",
                'Media Type': placement.media_type,
                'Google Docket': placement.docket_url or "No Google docket",
                'Local Docket': f'=HYPERLINK("./dockets/{docket_filename}", "Open Docket")',
                'Created': placement.created_at.strftime('%Y-%m-%d %H:%M:%S'),
                'Updated': placement.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
                'Notes': placement.notes or ""
            })
        
        # Create Excel file
        df = pd.DataFrame(data)
        excel_path = os.path.join(temp_dir, 'media_placements.xlsx')
        
        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Media Placements', index=False)
            
            # Auto-adjust columns' width
            worksheet = writer.sheets['Media Placements']
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.column_dimensions[chr(65 + idx)].width = max_length
        
        # Create README file
        readme_path = os.path.join(temp_dir, 'README.txt')
        with open(readme_path, 'w') as f:
            f.write("Media Placements Export Package\n")
            f.write("==============================\n\n")
            f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("Contents:\n")
            f.write("- media_placements.xlsx: Excel file with all media placements and links to dockets\n")
            f.write("- dockets/: Directory containing Word document dockets for each media placement\n\n")
            f.write("Instructions:\n")
            f.write("1. Keep the Excel file and dockets folder in the same directory\n")
            f.write("2. Open the Excel file to view all placements\n")
            f.write("3. Click on 'Open Docket' links to open the corresponding docket files\n")
        
        # Create ZIP file in memory
        memory_file = io.BytesIO()
        
        with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add Excel file
            zipf.write(excel_path, arcname='media_placements.xlsx')
            
            # Add README
            zipf.write(readme_path, arcname='README.txt')
            
            # Add all docket files
            for root, dirs, files in os.walk(dockets_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.join('dockets', file)
                    zipf.write(file_path, arcname=arcname)
        
        # Clean up temporary directory
        shutil.rmtree(temp_dir)
        
        # Prepare ZIP for download
        memory_file.seek(0)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return send_file(
            memory_file,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'media_placements_complete_{timestamp}.zip'
        )
        
    except Exception as e:
        app.logger.error(f"Error creating complete export package: {str(e)}")
        flash(f'Error creating export package: {str(e)}', 'danger')
        return redirect(url_for('dashboard'))

def create_docket_for_export(placement, output_path):
    """Create a Word docket for a specific placement and save to the given path."""
    try:
        # Create a new Word document
        doc = Document()
        
        # Add title with formatting
        doc.add_heading(placement.title or "Untitled Article", level=1)
        
        # Add basic info section
        doc.add_heading("Basic Information", level=2)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Add basic info rows
        cells = table.rows[0].cells
        cells[0].text = "URL"
        cells[1].text = placement.url
        
        cells = table.rows[1].cells
        cells[0].text = "Source"
        cells[1].text = placement.source or "Unknown"
        
        cells = table.rows[2].cells
        cells[0].text = "Publication Date"
        cells[1].text = str(placement.publication_date) if placement.publication_date else "Unknown"
        
        cells = table.rows[3].cells
        cells[0].text = "Media Type"
        cells[1].text = placement.media_type.title()
        
        cells = table.rows[4].cells
        cells[0].text = "Created"
        cells[1].text = placement.created_at.strftime('%Y-%m-%d %H:%M:%S')
        
        # Add screenshot section
        doc.add_heading("Screenshot", level=2)
        doc.add_paragraph("Screenshot attempted during export. If missing, visit the URL directly.")
        
        # Try to take a screenshot (with faster timeout)
        try:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_screenshot_path = temp_file.name
            
            # Take the screenshot with reduced timeout
            screenshot_path = take_screenshot(placement.url, temp_screenshot_path, timeout=10)
            
            if screenshot_path:
                # Add the screenshot to the document
                doc.add_picture(screenshot_path, width=Inches(6.0))
                doc.add_paragraph(f"Screenshot taken on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                # Clean up temp file
                try:
                    os.unlink(temp_screenshot_path)
                except:
                    pass
        except Exception as e:
            app.logger.error(f"Error taking screenshot for export: {str(e)}")
            doc.add_paragraph(f"Error capturing screenshot - please visit the URL directly.")
        
        # Add summary section
        doc.add_heading("Summary", level=2)
        
        # Try to extract a summary with shorter timeout
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(placement.url, headers=headers, timeout=5)
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Try to get the main article content
            article = soup.find('article') or soup.find(class_=['article', 'post', 'content', 'main-content'])
            
            if article:
                paragraphs = article.find_all('p')
                content = ' '.join([p.get_text().strip() for p in paragraphs[:3]])  # Get first 3 paragraphs
            else:
                # Fallback to all paragraphs with limit
                paragraphs = soup.find_all('p')[:5]
                content = ' '.join([p.get_text().strip() for p in paragraphs])
            
            # Limit summary length
            if content:
                content = content[:800] + '...' if len(content) > 800 else content
                doc.add_paragraph(content)
            else:
                doc.add_paragraph("No text content could be extracted from this page.")
                
        except Exception as e:
            app.logger.error(f"Error extracting summary for export: {str(e)}")
            doc.add_paragraph("Could not extract summary - please visit the URL directly.")
        
        # Add notes section
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(placement.notes or "No notes available")
        
        # Save the document to the specified path
        doc.save(output_path)
        return True
        
    except Exception as e:
        app.logger.error(f"Error creating docket for export: {str(e)}")
        return False

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('error.html', error='An internal error occurred'), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(host='0.0.0.0', port=5000, debug=True)
